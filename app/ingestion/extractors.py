from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from app.models.records import ExtractedDocument, ExtractedSection, SourceDocument

HEADING_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十百0-9]+[章节部分篇条]"),
    re.compile(r"^[0-9]+(\.[0-9]+)*[、.]"),
    re.compile(r"^[一二三四五六七八九十]+[、.]"),
    re.compile(r"^[A-Za-z][0-9]+(\.[0-9]+)*"),
]


def is_numeric_marker(text: str) -> bool:
    return bool(re.fullmatch(r"[0-9]+", text.strip()))


def normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def looks_like_heading(text: str) -> bool:
    if not text:
        return False
    compact = text.strip()
    if len(compact) > 60:
        return False
    if compact.endswith("：") and len(compact) <= 40:
        return True
    return any(pattern.match(compact) for pattern in HEADING_PATTERNS)


def is_toc_fragment(title: str, content: str) -> bool:
    title = title.strip()
    if "目录" in title:
        return True
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if not lines:
        return False
    toc_like = 0
    for line in lines:
        if re.search(r"\.{2,}\s*\d+$", line) or re.search(r"[^\s]\s+\d+$", line):
            toc_like += 1
    return toc_like >= max(1, len(lines) - 1)


def build_section(section_id: str, title: str, text_lines: list[str], order: int, page_hint: str | None = None) -> ExtractedSection | None:
    content = normalize_text("\n".join(line for line in text_lines if line.strip()))
    if not content:
        return None
    if is_toc_fragment(title, content):
        return None
    return ExtractedSection(
        section_id=section_id,
        title=title.strip() or "未命名章节",
        text=content,
        order=order,
        page_hint=page_hint,
        section_path=[title.strip() or "未命名章节"],
    )


def derive_entry_title(source_title: str, text: str, fallback: str) -> str:
    compact = normalize_text(text)
    if not compact:
        return fallback
    for separator in ("：", ":", "（", "("):
        head = compact.split(separator, 1)[0].strip(" -")
        if 2 <= len(head) <= 36:
            return head
    if len(compact) <= 28:
        return compact
    return f"{source_title} - {fallback}"


def build_tabular_sections(
    source: SourceDocument,
    rows: list[list[str]],
    order_start: int,
    page_hint: str | None = None,
) -> tuple[list[ExtractedSection], int]:
    sections: list[ExtractedSection] = []
    order = order_start
    current_major = ""
    current_minor = ""

    for row_index, row in enumerate(rows, start=1):
        cells = [normalize_text(cell) for cell in row if normalize_text(cell)]
        if not cells:
            continue
        if len(set(cells)) == 1 and "汇报人" in cells[0]:
            continue
        if cells[0] == "序号":
            current_major = cells[1] if len(cells) > 1 else current_major
            current_minor = ""
            continue

        title = source.title
        text_lines: list[str] = []

        if len(cells) >= 3 and is_numeric_marker(cells[0]):
            current_major = cells[1]
            current_minor = cells[2]
            title = " / ".join(part for part in [current_major, current_minor] if part) or f"{source.title} 条目 {cells[0]}"
            text_lines = cells[3:] or cells[1:]
        elif len(cells) == 2 and is_numeric_marker(cells[0]):
            title = derive_entry_title(source.title, cells[1], f"条目 {cells[0]}")
            text_lines = [cells[1]]
        elif len(cells) >= 2:
            label = cells[0]
            prefix = " / ".join(part for part in [current_major, label] if part and part != current_major)
            title = prefix or label
            text_lines = cells[1:]
        else:
            title = derive_entry_title(source.title, cells[0], f"条目 {row_index}")
            text_lines = cells

        section = build_section(
            section_id=f"{source.doc_id}-tabular-{order:04d}",
            title=title,
            text_lines=text_lines,
            order=order,
            page_hint=page_hint,
        )
        if section is not None:
            sections.append(section)
            order += 1

    return sections, order


def extract_docx(source: SourceDocument) -> ExtractedDocument:
    document = Document(source.file_path)
    sections: list[ExtractedSection] = []
    current_title = source.title
    current_lines: list[str] = []
    order = 0

    def flush() -> None:
        nonlocal order, current_lines
        section = build_section(
            section_id=f"{source.doc_id}-section-{order:04d}",
            title=current_title,
            text_lines=current_lines,
            order=order,
        )
        if section is not None:
            sections.append(section)
            order += 1
        current_lines = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        is_heading = style_name.lower().startswith("heading") or looks_like_heading(text)
        if is_heading:
            if current_lines:
                flush()
            current_title = text
            continue
        current_lines.append(text)

    if current_lines:
        flush()

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if cells:
                rows.append(cells)
        numeric_rows = sum(1 for cells in rows if cells and is_numeric_marker(cells[0]))
        if len(rows) >= 6 and numeric_rows >= max(2, len(rows) // 4):
            tabular_sections, order = build_tabular_sections(
                source=source,
                rows=rows,
                order_start=order,
                page_hint=f"table-{table_index}",
            )
            sections.extend(tabular_sections)
            continue
        row_lines = [" | ".join(cells) for cells in rows]
        section = build_section(
            section_id=f"{source.doc_id}-table-{table_index:04d}",
            title=f"{source.title} 表格 {table_index}",
            text_lines=row_lines,
            order=order,
        )
        if section is not None:
            sections.append(section)
            order += 1

    if not sections:
        sections.append(
            ExtractedSection(
                section_id=f"{source.doc_id}-section-0000",
                title=source.title,
                text="",
                order=0,
                page_hint=None,
                section_path=[source.title],
            )
        )

    return ExtractedDocument(source=source, sections=sections)


def extract_txt(source: SourceDocument) -> ExtractedDocument:
    text = Path(source.file_path).read_text(encoding="utf-8", errors="ignore")
    sections: list[ExtractedSection] = []
    title = source.title
    buffer: list[str] = []
    order = 0
    for raw_line in text.splitlines():
        line = normalize_text(raw_line)
        if not line:
            continue
        if looks_like_heading(line):
            section = build_section(
                section_id=f"{source.doc_id}-section-{order:04d}",
                title=title,
                text_lines=buffer,
                order=order,
            )
            if section is not None:
                sections.append(section)
                order += 1
            title = line
            buffer = []
            continue
        buffer.append(line)

    section = build_section(
        section_id=f"{source.doc_id}-section-{order:04d}",
        title=title,
        text_lines=buffer,
        order=order,
    )
    if section is not None:
        sections.append(section)

    return ExtractedDocument(source=source, sections=sections)


def extract_xlsx(source: SourceDocument) -> ExtractedDocument:
    workbook = load_workbook(source.file_path, data_only=True)
    sections: list[ExtractedSection] = []
    order = 0
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            values = [normalize_text(str(value)) for value in row if value is not None and normalize_text(str(value))]
            if values:
                rows.append(values)
        if len(rows) >= 2:
            tabular_sections, order = build_tabular_sections(
                source=source,
                rows=rows,
                order_start=order,
                page_hint=sheet_name,
            )
            sections.extend(tabular_sections)
            continue
        row_lines = [" | ".join(values) for values in rows]
        section = build_section(
            section_id=f"{source.doc_id}-sheet-{order:04d}",
            title=f"{source.title} - {sheet_name}",
            text_lines=row_lines,
            order=order,
            page_hint=sheet_name,
        )
        if section is not None:
            sections.append(section)
            order += 1
    return ExtractedDocument(source=source, sections=sections)


def extract_pdf(source: SourceDocument) -> ExtractedDocument:
    try:
        completed = subprocess.run(
            ["pdftotext", source.file_path, "-"],
            check=True,
            capture_output=True,
            text=True,
        )
    except (FileNotFoundError, subprocess.CalledProcessError) as error:
        raise RuntimeError(f"pdftotext failed for {source.file_path}") from error

    raw_text = completed.stdout
    pages = raw_text.split("\f")
    sections: list[ExtractedSection] = []
    order = 0
    for page_index, page_text in enumerate(pages, start=1):
        current_title = f"{source.title} - 第{page_index}页"
        buffer: list[str] = []
        for raw_line in page_text.splitlines():
            line = normalize_text(raw_line)
            if not line:
                continue
            if looks_like_heading(line):
                section = build_section(
                    section_id=f"{source.doc_id}-section-{order:04d}",
                    title=current_title,
                    text_lines=buffer,
                    order=order,
                    page_hint=f"p{page_index}",
                )
                if section is not None:
                    sections.append(section)
                    order += 1
                current_title = line
                buffer = []
                continue
            buffer.append(line)

        section = build_section(
            section_id=f"{source.doc_id}-section-{order:04d}",
            title=current_title,
            text_lines=buffer,
            order=order,
            page_hint=f"p{page_index}",
        )
        if section is not None:
            sections.append(section)
            order += 1

    return ExtractedDocument(source=source, sections=sections)


def extract_document(source: SourceDocument) -> ExtractedDocument:
    file_type = source.file_type.lower()
    if file_type == "docx":
        return extract_docx(source)
    if file_type == "pdf":
        return extract_pdf(source)
    if file_type == "xlsx":
        return extract_xlsx(source)
    if file_type == "txt":
        return extract_txt(source)
    raise ValueError(f"Unsupported file type: {source.file_type}")
